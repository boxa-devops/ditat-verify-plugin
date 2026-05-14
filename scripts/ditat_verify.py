#!/usr/bin/env python3
"""
Ditat shipment verification helper — drives the /ditat-verify Claude skill.

Subcommands:
  fetch      — list unprocessed shipments, fetch details + download docs,
               emit JSON descriptor on stdout for Claude to consume.
  verify-one — fetch + download for one specific shipment (re-verify path).
  mark       — record a shipment as processed (after Claude wrote its report).
  status     — show processed shipments.
  reset      — clear processed flag for a shipment.
  check-env  — validate .env / credentials without making API calls.

Logs go to STDERR. Subcommand JSON output goes to STDOUT — safe to pipe.

State stored in state.db (SQLite, project root). Schema:
  processed_shipments(shipment_key PK, shipment_id, processed_at, report_path,
                      critical_count, warn_count)

Reuses ditat/client.py, ditat/services.py, ditat/extractors.py as-is.
"""

import argparse
import json
import logging
import os
import sqlite3
import sys
from datetime import datetime, timedelta, timezone
from pathlib import Path

from ditat.client import DitatClient
from ditat.config import Config
from ditat.extractors import KeyExtractor
from ditat.logging_utils import setup_logging
from ditat.response import ResponseParser
from ditat.services import DocumentService, ShipmentService

log = logging.getLogger("ditat")


def _state_root() -> Path:
    """State directory: env override → CLAUDE_PROJECT_DIR → cwd."""
    return Path(
        os.getenv("DITAT_STATE_DIR")
        or os.getenv("CLAUDE_PROJECT_DIR")
        or Path.cwd()
    ).resolve()


STATE_DB = _state_root() / os.getenv("DITAT_STATE_DB_NAME", "state.db")
REPORTS_DIR = Path(os.getenv("DITAT_REPORTS_DIR") or (_state_root() / "reports"))


def db_connect() -> sqlite3.Connection:
    conn = sqlite3.connect(STATE_DB)
    conn.execute("PRAGMA journal_mode=WAL")
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS processed_shipments (
            shipment_key   TEXT PRIMARY KEY,
            shipment_id    TEXT,
            processed_at   TEXT NOT NULL,
            report_path    TEXT,
            critical_count INTEGER DEFAULT 0,
            warn_count     INTEGER DEFAULT 0
        )
        """
    )
    return conn


def _ditat_summary(details: dict) -> dict:
    """Flatten the most useful shipment fields for comparison against docs."""
    eg = ResponseParser._ci_get(details, "EntityGraph") or details or {}
    if not isinstance(eg, dict):
        return {}
    pick = lambda *keys: ResponseParser._ci_get(eg, *keys)
    stops = pick("dspShipmentStops") or []
    items = pick("dspShipmentItems") or []
    revenues = pick("rnpShipmentRevenues") or []

    def stop_summary(stop_type_substring: str) -> dict:
        if not isinstance(stops, list):
            return {}
        for s in stops:
            if not isinstance(s, dict):
                continue
            stype = str(ResponseParser._ci_get(s, "stopType", "type") or "").lower()
            if stop_type_substring in stype:
                return {
                    "name":     ResponseParser._ci_get(s, "name", "locationName"),
                    "city":     ResponseParser._ci_get(s, "city"),
                    "state":    ResponseParser._ci_get(s, "state"),
                    "zip":      ResponseParser._ci_get(s, "zip", "zipCode"),
                    "appointment_from": ResponseParser._ci_get(s, "appointmentFrom", "scheduledFrom"),
                    "appointment_to":   ResponseParser._ci_get(s, "appointmentTo", "scheduledTo"),
                    "arrived":   ResponseParser._ci_get(s, "arrived", "arrivedOn"),
                    "departed":  ResponseParser._ci_get(s, "departed", "departedOn"),
                    "stop_type": stype,
                }
        return {}

    total_weight = 0.0
    total_pieces = 0
    commodities = []
    for it in items if isinstance(items, list) else []:
        if not isinstance(it, dict):
            continue
        w = ResponseParser._ci_get(it, "weight", "weightLbs")
        p = ResponseParser._ci_get(it, "pieces", "quantity")
        c = ResponseParser._ci_get(it, "commodity", "description")
        try:
            total_weight += float(w) if w is not None else 0.0
        except (TypeError, ValueError):
            pass
        try:
            total_pieces += int(p) if p is not None else 0
        except (TypeError, ValueError):
            pass
        if c:
            commodities.append(c)

    total_revenue = 0.0
    for r in revenues if isinstance(revenues, list) else []:
        if not isinstance(r, dict):
            continue
        amt = ResponseParser._ci_get(r, "amount", "total", "totalAmount")
        try:
            total_revenue += float(amt) if amt is not None else 0.0
        except (TypeError, ValueError):
            pass

    return {
        "shipment_id":    pick("id", "shipmentId"),
        "shipment_key":   pick("key", "shipmentKey"),
        "status":         pick("status"),
        "bol_number":     pick("bolNumber", "bol"),
        "load_number":    pick("loadNumber", "loadId"),
        "po_numbers":     pick("poNumbers", "poNumber", "purchaseOrders"),
        "reference":      pick("referenceNumber", "reference"),
        "customer":       pick("customer", "customerName", "billTo"),
        "carrier":        pick("carrier", "carrierName"),
        "equipment_type": pick("equipmentType", "trailerType"),
        "hazmat":         pick("hazmat", "isHazmat"),
        "total_weight_lbs": total_weight or None,
        "total_pieces":   total_pieces or None,
        "commodities":    commodities or None,
        "total_revenue":  total_revenue or None,
        "pickup":         stop_summary("pickup") or stop_summary("origin"),
        "delivery":       stop_summary("delivery") or stop_summary("destination") or stop_summary("consignee"),
        "billing_notes":  pick("billingNotes"),
    }


def _build_services(config: Config):
    client = DitatClient(config.base_url, config.account_id, config.client_id, config.client_secret)
    return client, ShipmentService(client), DocumentService(client), KeyExtractor()


def _process_one(shipment_obj: dict, details: dict, config: Config,
                 document_svc: DocumentService, key_extractor: KeyExtractor) -> dict | None:
    """Build per-shipment JSON record. Returns None if no usable key."""
    key = key_extractor.extract(shipment_obj, "shipment")
    if not key:
        return None
    docs = document_svc.list_documents(details)
    out_dir = config.download_dir / str(key)
    doc_records = []
    for d in docs:
        path = document_svc.download(d, shipment_key=key, out_dir=out_dir)
        if not path:
            continue
        doc_records.append({
            "doc_key":   key_extractor.extract(d, "document"),
            "file_name": d.get("fileName") or d.get("FileName") or path.name,
            "file_type": d.get("fileType") or d.get("FileType"),
            "ordinal":   d.get("ordinal"),
            "path":      str(path.resolve()),
        })
    return {
        "shipment_key": key,
        "shipment_id":  ResponseParser._ci_get(shipment_obj, "id", "shipmentId"),
        "ditat_fields": _ditat_summary(details),
        "documents":    doc_records,
    }


def cmd_check_env(_args: argparse.Namespace) -> int:
    config = Config()
    ok, missing = config.validate()
    print(json.dumps({
        "ok": ok,
        "missing": missing,
        "base_url": config.base_url,
        "download_dir": str(config.download_dir.resolve()),
        "state_db": str(STATE_DB.resolve()),
    }, indent=2))
    return 0 if ok else 2


def cmd_fetch(args: argparse.Namespace) -> int:
    config = Config()
    ok, missing = config.validate()
    if not ok:
        print(json.dumps({"error": f"missing env vars: {missing}"}), file=sys.stderr)
        return 2

    _, shipment_svc, document_svc, key_extractor = _build_services(config)

    since = None if args.all else datetime.now(timezone.utc) - timedelta(days=args.since_days)
    shipments = shipment_svc.list_shipments(since=since, filter_column=args.filter_column)

    conn = db_connect()
    processed_keys = {row[0] for row in conn.execute("SELECT shipment_key FROM processed_shipments")}
    conn.close()

    out: list[dict] = []
    seen = 0
    for s in shipments:
        if seen >= args.limit:
            break
        key = key_extractor.extract(s, "shipment")
        if not key:
            continue
        if not args.include_processed and key in processed_keys:
            continue
        try:
            details = shipment_svc.get_details(key)
        except Exception as e:
            log.error("detail fetch failed for %s: %s", key, e)
            continue
        rec = _process_one(s, details, config, document_svc, key_extractor)
        if not rec:
            continue
        if args.require_docs and not rec["documents"]:
            continue
        out.append(rec)
        seen += 1

    print(json.dumps({"count": len(out), "shipments": out}, default=str, indent=2))
    return 0


def cmd_verify_one(args: argparse.Namespace) -> int:
    config = Config()
    ok, missing = config.validate()
    if not ok:
        print(json.dumps({"error": f"missing env vars: {missing}"}), file=sys.stderr)
        return 2

    _, shipment_svc, document_svc, key_extractor = _build_services(config)

    try:
        details = shipment_svc.get_details(args.shipment_key)
    except Exception as e:
        print(json.dumps({"error": f"detail fetch failed: {e}"}), file=sys.stderr)
        return 1

    stub = {"Key": args.shipment_key}
    rec = _process_one(stub, details, config, document_svc, key_extractor)
    if not rec:
        print(json.dumps({"error": "could not extract shipment key"}), file=sys.stderr)
        return 1

    print(json.dumps({"count": 1, "shipments": [rec]}, default=str, indent=2))
    return 0


def cmd_mark(args: argparse.Namespace) -> int:
    conn = db_connect()
    conn.execute(
        "INSERT OR REPLACE INTO processed_shipments "
        "(shipment_key, shipment_id, processed_at, report_path, critical_count, warn_count) "
        "VALUES (?, ?, ?, ?, ?, ?)",
        (
            args.shipment_key,
            args.shipment_id,
            datetime.now(timezone.utc).isoformat(),
            args.report_path,
            args.critical,
            args.warn,
        ),
    )
    conn.commit()
    conn.close()
    print(json.dumps({"marked": args.shipment_key}))
    return 0


def cmd_status(_args: argparse.Namespace) -> int:
    conn = db_connect()
    cur = conn.execute(
        "SELECT shipment_key, shipment_id, processed_at, critical_count, warn_count, report_path "
        "FROM processed_shipments ORDER BY processed_at DESC LIMIT 20"
    )
    rows = [dict(zip([c[0] for c in cur.description], r)) for r in cur.fetchall()]
    total = conn.execute("SELECT COUNT(*) FROM processed_shipments").fetchone()[0]
    conn.close()
    print(json.dumps({"total_processed": total, "recent": rows}, indent=2))
    return 0


def cmd_reset(args: argparse.Namespace) -> int:
    conn = db_connect()
    conn.execute("DELETE FROM processed_shipments WHERE shipment_key=?", (args.shipment_key,))
    conn.commit()
    conn.close()
    print(json.dumps({"reset": args.shipment_key}))
    return 0


def _select_report_paths(args: argparse.Namespace) -> list[Path]:
    """Pick which .md reports to bundle into the docx."""
    if args.keys:
        keys = [k.strip() for k in args.keys.split(",") if k.strip()]
        return [REPORTS_DIR / f"{k}.md" for k in keys]

    conn = db_connect()
    if args.since_days is not None:
        cutoff = (datetime.now(timezone.utc) - timedelta(days=args.since_days)).isoformat()
        cur = conn.execute(
            "SELECT shipment_key, report_path FROM processed_shipments "
            "WHERE processed_at >= ? ORDER BY processed_at",
            (cutoff,),
        )
    else:
        cur = conn.execute(
            "SELECT shipment_key, report_path FROM processed_shipments ORDER BY processed_at"
        )
    rows = cur.fetchall()
    conn.close()

    paths: list[Path] = []
    for key, rpath in rows:
        if rpath:
            p = Path(rpath)
            if not p.is_absolute():
                p = _state_root() / p
        else:
            p = REPORTS_DIR / f"{key}.md"
        if p.exists():
            paths.append(p)
    return paths


def _md_to_docx(paths: list[Path], out_path: Path) -> None:
    """Render markdown reports into one .docx. Minimal MD subset only (headings,
    lists, fenced code, bold). Keeps it dependency-light and stable."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.core_properties.title = "Ditat Verification Report"
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("Ditat Verification Report", level=0)
    doc.add_paragraph(
        f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}  "
        f"| Shipments: {len(paths)}"
    )

    for i, p in enumerate(paths):
        if i > 0:
            doc.add_page_break()
        _render_md_into(doc, p.read_text(encoding="utf-8"))

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def _render_md_into(doc, md: str) -> None:
    """Tiny markdown → docx renderer for our report format."""
    import re
    in_code = False
    code_buf: list[str] = []
    bold_re = re.compile(r"\*\*(.+?)\*\*")

    def flush_code():
        nonlocal code_buf
        if not code_buf:
            return
        p = doc.add_paragraph()
        run = p.add_run("\n".join(code_buf))
        run.font.name = "Consolas"
        from docx.shared import Pt
        run.font.size = Pt(9)
        code_buf = []

    def add_inline(par, text: str):
        pos = 0
        for m in bold_re.finditer(text):
            if m.start() > pos:
                par.add_run(text[pos:m.start()])
            r = par.add_run(m.group(1))
            r.bold = True
            pos = m.end()
        if pos < len(text):
            par.add_run(text[pos:])

    for raw in md.splitlines():
        line = raw.rstrip()
        if line.startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_buf.append(raw)
            continue
        if not line.strip():
            doc.add_paragraph("")
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.lstrip().startswith(("- ", "* ")):
            indent = len(line) - len(line.lstrip())
            text = line.lstrip()[2:]
            par = doc.add_paragraph(style="List Bullet")
            if indent >= 2:
                par.paragraph_format.left_indent = None
            add_inline(par, text)
        else:
            par = doc.add_paragraph()
            add_inline(par, line)
    flush_code()


def cmd_build_docx(args: argparse.Namespace) -> int:
    paths = _select_report_paths(args)
    if not paths:
        print(json.dumps({"error": "no reports matched"}), file=sys.stderr)
        return 1

    if args.output:
        out = Path(args.output)
        if not out.is_absolute():
            out = _state_root() / out
    else:
        stamp = datetime.now(timezone.utc).strftime("%Y-%m-%d")
        out = REPORTS_DIR / f"ditat-verify-{stamp}.docx"

    try:
        _md_to_docx(paths, out)
    except ImportError:
        print(json.dumps({
            "error": "python-docx not installed. Run: pip install -r scripts/requirements.txt"
        }), file=sys.stderr)
        return 2

    print(json.dumps({
        "docx": str(out.resolve()),
        "shipments": len(paths),
        "sources": [str(p) for p in paths],
    }, indent=2))
    return 0


def main() -> int:
    p = argparse.ArgumentParser(description="Ditat shipment verification helper")
    p.add_argument("--verbose", action="store_true")
    sub = p.add_subparsers(dest="cmd", required=True)

    f = sub.add_parser("fetch", help="Emit JSON for unprocessed shipments + download docs")
    f.add_argument("--limit", type=int, default=5)
    f.add_argument("--since-days", type=int, default=30)
    f.add_argument("--all", action="store_true", help="Ignore since-days, list all shipments")
    f.add_argument("--filter-column", default="updatedOn")
    f.add_argument("--include-processed", action="store_true",
                   help="Include shipments already in state.db (re-verify)")
    f.add_argument("--require-docs", action="store_true", default=True,
                   help="Skip shipments with zero downloadable docs")
    f.set_defaults(func=cmd_fetch)

    m = sub.add_parser("mark", help="Mark a shipment as processed")
    m.add_argument("shipment_key")
    m.add_argument("--shipment-id", default=None)
    m.add_argument("--report-path", default=None)
    m.add_argument("--critical", type=int, default=0)
    m.add_argument("--warn", type=int, default=0)
    m.set_defaults(func=cmd_mark)

    s = sub.add_parser("status", help="Show processed-shipment status")
    s.set_defaults(func=cmd_status)

    r = sub.add_parser("reset", help="Clear processed flag for a shipment")
    r.add_argument("shipment_key")
    r.set_defaults(func=cmd_reset)

    v = sub.add_parser("verify-one", help="Fetch + download docs for one specific shipment")
    v.add_argument("shipment_key")
    v.set_defaults(func=cmd_verify_one)

    e = sub.add_parser("check-env", help="Validate .env / credentials without API calls")
    e.set_defaults(func=cmd_check_env)

    b = sub.add_parser("build-docx", help="Bundle markdown reports into one .docx")
    b.add_argument("--since-days", type=int, default=None,
                   help="Include reports processed in the last N days (default: all)")
    b.add_argument("--keys", default=None,
                   help="Comma-separated shipment_keys to include (overrides --since-days)")
    b.add_argument("--output", default=None,
                   help="Output .docx path (default: reports/ditat-verify-<date>.docx)")
    b.set_defaults(func=cmd_build_docx)

    args = p.parse_args()
    setup_logging(args.verbose)
    REPORTS_DIR.mkdir(exist_ok=True)
    try:
        return args.func(args)
    except Exception as e:
        log.exception("helper failed")
        print(json.dumps({"error": str(e)}), file=sys.stderr)
        return 1


if __name__ == "__main__":
    sys.exit(main())
