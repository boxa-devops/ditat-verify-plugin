"""Build the single batch .docx report.

Renders straight from in-memory data — no per-shipment markdown step. Output:
  1. Title + generation timestamp + batch counters
  2. Summary table: every shipment in the batch (verdict + counts)
  3. Detail section: ONLY problematic shipments (critical, warn, or RC missing)

`build_batch_docx(batch, findings_index, diff_index, out_path)` is the entrypoint.

  batch          : the .ditat_batch.json content (slim Ditat records keyed by
                   shipment_key, plus shipment_id and document list)
  findings_index : { shipment_key -> agent-extracted dict { rc, bol, pod, ... } }
  diff_index     : { shipment_key -> diff.run_diff() result }
"""

from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, RGBColor

from .classify import doc_present


SEVERITY_ORDER = {"ISSUES": 0, "RC MISSING": 1, "WARN": 2, "OK": 3}

SEVERITY_COLOR = {
    "ISSUES":     RGBColor(0xB0, 0x00, 0x20),  # red
    "RC MISSING": RGBColor(0xA0, 0x52, 0x00),  # amber
    "WARN":       RGBColor(0xA0, 0x52, 0x00),
    "OK":         RGBColor(0x10, 0x70, 0x20),  # green
}


def _init_doc() -> Document:
    doc = Document()
    doc.core_properties.title = "Ditat Verification Report"
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    return doc


def _add_summary_table(doc: Document, rows: list[dict]) -> None:
    doc.add_heading("Summary", level=1)
    if not rows:
        doc.add_paragraph("No shipments in this batch.")
        return

    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, label in enumerate(["Shipment ID", "Key", "Verdict", "Critical", "Warn", "Docs"]):
        hdr[i].text = label
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True

    rows_sorted = sorted(
        rows,
        key=lambda r: (SEVERITY_ORDER.get(r["verdict"], 9), r.get("shipment_key", "")),
    )
    for r in rows_sorted:
        cells = table.add_row().cells
        cells[0].text = str(r.get("shipment_id") or "")
        cells[1].text = str(r.get("shipment_key") or "")
        v_cell = cells[2]
        v_cell.text = r["verdict"]
        color = SEVERITY_COLOR.get(r["verdict"])
        if color is not None:
            for p in v_cell.paragraphs:
                for run in p.runs:
                    run.font.color.rgb = color
                    run.bold = True
        cells[3].text = str(r.get("critical_count", 0))
        cells[4].text = str(r.get("warn_count", 0))
        cells[5].text = r.get("docs_label", "")


# Per-finding severity colors (distinct from the verdict colors above).
FINDING_SEV_COLOR = {
    "critical":   RGBColor(0xB0, 0x00, 0x20),  # red
    "warn":       RGBColor(0xA0, 0x52, 0x00),  # amber
    "info":       RGBColor(0x70, 0x70, 0x70),  # gray
    "RC MISSING": RGBColor(0xA0, 0x52, 0x00),  # amber
}

_FINDINGS_COLS = ["Shipment", "Severity", "Pair", "Field", "Value", "vs RC", "Note"]


def _add_findings_table(doc: Document, rows: list[dict]) -> None:
    """One consolidated table — every finding across the batch is a row."""
    table = doc.add_table(rows=1, cols=len(_FINDINGS_COLS))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, label in enumerate(_FINDINGS_COLS):
        hdr[i].text = label
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True

    for r in rows:
        cells = table.add_row().cells
        cells[0].text = str(r["shipment"])
        sev_cell = cells[1]
        sev_cell.text = r["severity"]
        color = FINDING_SEV_COLOR.get(r["severity"])
        if color is not None:
            for p in sev_cell.paragraphs:
                for run in p.runs:
                    run.font.color.rgb = color
                    run.bold = True
        cells[2].text = str(r["pair"])
        cells[3].text = str(r["field"])
        cells[4].text = str(r["value"])
        cells[5].text = str(r["vs"])
        cells[6].text = str(r["note"])


_SEV_RANK = {"critical": 0, "warn": 1, "info": 2, "RC MISSING": 3}


def _findings_rows(problematic_keys, batch, diff_index, findings_index) -> list[dict]:
    """Flatten every problematic shipment's findings into table rows."""
    by_key = {e.get("shipment_key"): e for e in batch}
    rows: list[dict] = []
    for key in problematic_keys:
        entry = by_key.get(key, {})
        ship = entry.get("shipment_id") or key
        diff_result = diff_index.get(key) or {}
        verdict = diff_result.get("verdict", "RC MISSING")
        crit = diff_result.get("critical") or []
        warn = diff_result.get("warn") or []
        if crit or warn:
            for f in [*crit, *warn]:
                rows.append({
                    "shipment": ship, "severity": f["severity"],
                    "pair": f["pair"], "field": f["field"],
                    "value": f["a"], "vs": f["b"], "note": f["message"],
                })
        else:
            rows.append({
                "shipment": ship, "severity": verdict, "pair": "—", "field": "—",
                "value": "", "vs": "", "note": "RC missing — no cross-check",
            })
        missing = (findings_index.get(key) or {}).get("docs_missing") or []
        if missing:
            rows.append({
                "shipment": ship, "severity": "info", "pair": "docs",
                "field": "missing", "value": "", "vs": ", ".join(missing),
                "note": "not provided",
            })
    rows.sort(key=lambda r: (str(r["shipment"]), _SEV_RANK.get(r["severity"], 9)))
    return rows


def build_batch_docx(
    batch: list[dict],
    findings_index: dict,
    diff_index: dict,
    out_path: Path,
    anomalies_only: bool = True,
) -> dict:
    """Render the final docx. Returns counters for the helper to print.

    anomalies_only=True (default) drops the per-shipment summary table —
    docx contains only the counts header + detail sections for problematic
    shipments. Set False to include the full summary table.
    """
    doc = _init_doc()
    doc.add_heading("Ditat Verification Report", level=0)

    now = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")

    summary_rows: list[dict] = []
    problematic_keys: list[str] = []
    counts = {"OK": 0, "WARN": 0, "ISSUES": 0, "RC MISSING": 0}
    for entry in batch:
        key = entry.get("shipment_key")
        diff_result = diff_index.get(key) or {"verdict": "RC MISSING",
                                              "critical_count": 0, "warn_count": 0}
        verdict = diff_result.get("verdict", "RC MISSING")
        counts[verdict] = counts.get(verdict, 0) + 1
        docs = entry.get("documents") or []
        doc_marks = [f"{lbl}{'✓' if doc_present(docs, lbl) else '✗'}"
                     for lbl in ("RC", "BOL", "POD")]
        summary_rows.append({
            "shipment_key": key,
            "shipment_id": entry.get("shipment_id"),
            "verdict": verdict,
            "critical_count": diff_result.get("critical_count", 0),
            "warn_count": diff_result.get("warn_count", 0),
            "docs_label": " · ".join(doc_marks),
        })
        if verdict in {"ISSUES", "WARN", "RC MISSING"}:
            problematic_keys.append(key)

    intro = doc.add_paragraph()
    intro.add_run(f"Generated: {now}    ").italic = True
    intro.add_run(f"Shipments: {len(batch)}    ").italic = True
    intro.add_run(
        f"OK: {counts.get('OK', 0)}  ·  WARN: {counts.get('WARN', 0)}  ·  "
        f"ISSUES: {counts.get('ISSUES', 0)}  ·  RC MISSING: {counts.get('RC MISSING', 0)}"
    ).italic = True

    if not anomalies_only:
        _add_summary_table(doc, summary_rows)

    if problematic_keys:
        if not anomalies_only:
            doc.add_page_break()
        doc.add_heading("Findings", level=1)
        rows = _findings_rows(problematic_keys, batch, diff_index, findings_index)
        _add_findings_table(doc, rows)
    else:
        doc.add_paragraph("")
        p = doc.add_paragraph()
        p.add_run("All shipments passed. No problematic shipments in this batch.").italic = True

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return {
        "shipments": len(batch),
        "problematic": len(problematic_keys),
        "verdicts": counts,
    }
